#!/usr/bin/env python3
"""Excel Image Reviewer - Convert Excel sheets to images and review with LLM vision."""

import argparse
import base64
import io
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

import pandas as pd
from dotenv import load_dotenv
from openai import OpenAI
from PIL import Image, ImageDraw, ImageFont
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches
import openpyxl
from openpyxl.utils import get_column_letter
import markdown

load_dotenv()


class ExcelImageReviewer:
    """Convert Excel sheets to images and review them using LLM vision."""

    def __init__(self, excel_path, output_dir="output", model_name=None, base_url=None):
        self.excel_path = excel_path
        self.output_dir = Path(output_dir)
        self.model_name = model_name or os.getenv("OPENAI_MODEL", "gpt-4o")
        self.output_dir.mkdir(exist_ok=True)
        self.sheet_images = {}
        self.sheet_reviews = {}
        client_kwargs = {"api_key": os.getenv("OPENAI_API_KEY")}
        resolved_url = base_url or os.getenv("OPENAI_BASE_URL")
        if resolved_url:
            client_kwargs["base_url"] = resolved_url
        self.client = OpenAI(**client_kwargs)

    def excel_to_image(self, sheet_name):
        """Convert an Excel sheet to a PIL Image covering only the populated content area."""
        wb = openpyxl.load_workbook(self.excel_path, data_only=True)
        ws = wb[sheet_name]

        if ws.max_row is None or ws.max_column is None:
            print(f"Warning: Sheet '{sheet_name}' is empty")
            return None

        min_r, max_r = ws.min_row, ws.max_row
        min_c, max_c = ws.min_column, ws.max_column

        # Trim trailing blank rows
        while max_r >= min_r:
            if any(ws.cell(max_r, c).value not in (None, "") for c in range(min_c, max_c + 1)):
                break
            max_r -= 1

        # Trim trailing blank columns
        while max_c >= min_c:
            if any(ws.cell(r, max_c).value not in (None, "") for r in range(min_r, max_r + 1)):
                break
            max_c -= 1

        if max_r < min_r or max_c < min_c:
            print(f"Warning: Sheet '{sheet_name}' is empty")
            return None

        rows = [
            [str(ws.cell(r, c).value) if ws.cell(r, c).value is not None else ""
             for c in range(min_c, max_c + 1)]
            for r in range(min_r, max_r + 1)
        ]

        col_widths = []
        for c in range(min_c, max_c + 1):
            dim = ws.column_dimensions.get(get_column_letter(c))
            w = dim.width if (dim and dim.width) else 10
            col_widths.append(max(80, int(w * 7)))

        row_heights = []
        for r in range(min_r, max_r + 1):
            dim = ws.row_dimensions.get(r)
            h = dim.height if (dim and dim.height) else 15
            row_heights.append(max(22, int(h * 1.33)))

        return self._render_table_image(rows, col_widths, row_heights, sheet_name)

    def _render_table_image(self, rows, col_widths, row_heights, title):
        """Render a 2D list of cell strings as a PIL Image table."""
        PADDING, TITLE_H = 10, 35
        total_w = PADDING * 2 + sum(col_widths)
        total_h = PADDING + TITLE_H + sum(row_heights) + PADDING

        img = Image.new("RGB", (total_w, total_h), "white")
        draw = ImageDraw.Draw(img)

        def load_font(candidates, size):
            for path in candidates:
                if os.path.exists(path):
                    try:
                        return ImageFont.truetype(path, size)
                    except Exception:
                        pass
            return ImageFont.load_default()

        bold_candidates = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/Supplemental/Songti.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
            "/Library/Fonts/Arial Bold.ttf",
        ]
        reg_candidates = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Supplemental/Songti.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/Library/Fonts/Arial.ttf",
        ]
        title_font = load_font(bold_candidates, 13)
        bold_font = load_font(bold_candidates, 11)
        cell_font = load_font(reg_candidates, 10)

        draw.text((PADDING, PADDING), title, fill="#2c3e50", font=title_font)
        y = PADDING + TITLE_H

        for row_idx, (row, rh) in enumerate(zip(rows, row_heights)):
            x = PADDING
            is_header = row_idx == 0
            for cell_val, cw in zip(row, col_widths):
                fill = "#cfe2f3" if is_header else ("white" if row_idx % 2 == 0 else "#f5f5f5")
                draw.rectangle([(x, y), (x + cw - 1, y + rh - 1)], fill=fill, outline="#bbbbbb")
                font = bold_font if is_header else cell_font
                max_chars = max(4, (cw - 8) // 6)
                text = cell_val[:max_chars] + ("\u2026" if len(cell_val) > max_chars else "")
                draw.text((x + 4, y + max(2, (rh - 12) // 2)), text, fill="#1a1a1a", font=font)
                x += cw
            y += rh

        return img

    def review_image(self, image, sheet_name):
        """Review a sheet image using LLM vision and return the review text."""
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        img_base64 = base64.b64encode(buffered.getvalue()).decode()

        prompt = (
            f"你是IT审计与数据质量双领域审阅专家。请审阅该电子表格截图（sheet: {sheet_name}），"
            "仅使用中文输出标准 Markdown。\n\n"
            "目标：直接给出问题，定位到具体语句并给出可执行修改，不要写笼统总结。\n\n"
            "请按以下固定结构输出：\n"
            "## 一、逐条问题定位与整改意见\n"
            "- 仅列出有问题项，最多10条，按高/中/低排序。\n"
            "- 每条必须包含：具体位置（如表名+单元格/行列描述）、原文、问题说明、修改为、优先级。\n"
            "- 必须指出具体哪句话有问题，不能只做概括。\n"
            "- 不要写影响，不要写背景。\n\n"
            "## 二、审计程序要求符合性\n"
            "- 判断测试过程描述是否满足审计程序要求（如：测试目标清晰、抽样依据、样本量与覆盖、执行步骤、证据链、结论对应）。\n"
            "- 仅列出“不符合/证据不足”项，每条必须包含：具体位置、原文、问题说明、修改为。\n\n"
            "## 三、测试问题的专业判断\n"
            "- 从IT控制测试专业角度指出方法性问题（如：设计有效与执行有效逻辑冲突、样本代表性不足、证据不可追溯、结论与记录不一致）。\n"
            "- 每条必须包含：具体位置、原文、专业问题、修改为（含责任角色+动作）。\n\n"
            "输出约束：\n"
            "- 不要复述表格内容。\n"
            "- 不输出空泛总结。\n"
            "- 不要写影响。\n"
            "- 对无法从截图确认的内容，标记为“需补充证据”。\n"
            "- 如未发现问题，仅输出“未发现需要整改的问题”。"
        )

        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {
                        "role": "system",
                        "content": "你是严谨的电子表格与数据质量审阅专家，请始终用中文 Markdown 作答。",
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_base64}"}},
                        ],
                    },
                ],
                max_tokens=2000,
                temperature=0.3,
                timeout=120,
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"Error reviewing image: {e}")
            return f"Error reviewing image: {e}"

    def _add_inline_runs(self, paragraph, node):
        """Append HTML inline nodes into a docx paragraph while preserving simple styles."""
        if isinstance(node, NavigableString):
            text = str(node)
            if text:
                paragraph.add_run(text)
            return
        if not isinstance(node, Tag):
            return

        if node.name == "br":
            paragraph.add_run("\n")
            return

        if node.name in ("strong", "b"):
            run = paragraph.add_run(node.get_text())
            run.bold = True
            return
        if node.name in ("em", "i"):
            run = paragraph.add_run(node.get_text())
            run.italic = True
            return
        if node.name == "code":
            run = paragraph.add_run(node.get_text())
            run.font.name = "Courier New"
            return
        if node.name == "a":
            text = node.get_text()
            href = node.get("href", "")
            paragraph.add_run(f"{text} ({href})" if href else text)
            return

        for child in node.children:
            self._add_inline_runs(paragraph, child)

    def _append_markdown_to_doc(self, doc, markdown_text):
        """Render markdown text into a Word document."""
        html = markdown.markdown(
            markdown_text or "",
            extensions=["fenced_code", "tables", "sane_lists"],
        )
        soup = BeautifulSoup(html, "html.parser")

        def render_block(tag):
            if tag.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = min(6, int(tag.name[1]))
                p = doc.add_heading(level=level)
                for child in tag.children:
                    self._add_inline_runs(p, child)
                return

            if tag.name == "p":
                p = doc.add_paragraph()
                for child in tag.children:
                    self._add_inline_runs(p, child)
                return

            if tag.name in ("ul", "ol"):
                style = "List Bullet" if tag.name == "ul" else "List Number"
                for li in tag.find_all("li", recursive=False):
                    p = doc.add_paragraph(style=style)
                    for child in li.children:
                        self._add_inline_runs(p, child)
                return

            if tag.name == "pre":
                p = doc.add_paragraph()
                run = p.add_run(tag.get_text())
                run.font.name = "Courier New"
                return

            if tag.name == "blockquote":
                p = doc.add_paragraph()
                p.style = "Intense Quote"
                p.add_run(tag.get_text())
                return

            if tag.name == "table":
                rows = tag.find_all("tr")
                if not rows:
                    return
                col_count = max(len(r.find_all(["th", "td"])) for r in rows)
                table = doc.add_table(rows=0, cols=col_count)
                table.style = "Table Grid"
                for r in rows:
                    cells = r.find_all(["th", "td"])
                    row_cells = table.add_row().cells
                    for idx, cell in enumerate(cells):
                        row_cells[idx].text = re.sub(r"\s+", " ", cell.get_text(" ", strip=True))
                return

            # Fallback: write text for unknown block tags
            text = tag.get_text(" ", strip=True)
            if text:
                doc.add_paragraph(text)

        for elem in soup.contents:
            if isinstance(elem, NavigableString):
                raw = str(elem).strip()
                if raw:
                    doc.add_paragraph(raw)
                continue
            if isinstance(elem, Tag):
                render_block(elem)

    def _sanitize_review_markdown(self, markdown_text):
        """Remove forbidden summary fields from model output before report rendering."""
        if not markdown_text:
            return markdown_text

        kept_lines = []
        for line in markdown_text.splitlines():
            stripped = line.strip()
            if re.match(r"^(?:[-*+]\s*)?(?:\d+\.\s*)?影响(?:[:：]|\b)", stripped):
                continue
            kept_lines.append(line)
        return "\n".join(kept_lines)

    def _excel_to_images_libreoffice(self, sheet_names):
        """Convert all sheets via LibreOffice headless → PDF → PIL Images.
        Returns {sheet_name: PIL.Image} mapped by sheet order."""
        with tempfile.TemporaryDirectory() as tmpdir:
            convert_modes = [
                # Keep each sheet on a single PDF page to prevent partial captures
                'pdf:calc_pdf_Export:{"SinglePageSheets":{"type":"boolean","value":"true"}}',
                "pdf",
            ]
            from pdf2image import convert_from_path
            last_error = "LibreOffice conversion failed"

            for convert_to in convert_modes:
                result = subprocess.run(
                    ["soffice", "--headless", "--convert-to", convert_to,
                     "--outdir", tmpdir, str(self.excel_path)],
                    capture_output=True, timeout=120,
                )
                if result.returncode != 0:
                    stderr = result.stderr.decode().strip()
                    last_error = stderr or f"LibreOffice conversion failed ({convert_to})"
                    continue

                pdf_path = Path(tmpdir) / (Path(self.excel_path).stem + ".pdf")
                if not pdf_path.exists():
                    last_error = f"Expected PDF not found: {pdf_path}"
                    continue

                pages = convert_from_path(pdf_path, dpi=150)
                if len(pages) != len(sheet_names):
                    last_error = (
                        f"PDF page count ({len(pages)}) does not match sheet count "
                        f"({len(sheet_names)}) for mode {convert_to}"
                    )
                    continue

                return {name: page for name, page in zip(sheet_names, pages)}

            raise RuntimeError(
                f"{last_error}. To avoid incomplete screenshots, falling back to the built-in renderer."
            )

    def process_excel(self, sheets=None, limit=None):
        """Convert all sheets to images and review each one."""
        print(f"Processing: {self.excel_path}")
        try:
            all_sheet_names = pd.ExcelFile(self.excel_path).sheet_names
        except Exception as e:
            print(f"Error reading Excel file: {e}")
            return
        print(f"Found {len(all_sheet_names)} sheet(s): {', '.join(all_sheet_names)}")
        if sheets:
            sheet_names = [s for s in all_sheet_names if s in sheets]
        else:
            sheet_names = all_sheet_names
        if limit:
            sheet_names = sheet_names[:limit]
        if len(sheet_names) < len(all_sheet_names):
            print(f"Processing {len(sheet_names)} sheet(s): {', '.join(sheet_names)}")

        # Try LibreOffice for high-fidelity rendering; fall back to built-in PIL renderer
        libreoffice_images = {}
        if shutil.which("soffice"):
            try:
                print("Converting with LibreOffice...")
                # LibreOffice export is workbook-wide; map pages with full sheet order first.
                libreoffice_images = self._excel_to_images_libreoffice(all_sheet_names)
                print("LibreOffice conversion successful.")
            except Exception as e:
                print(f"LibreOffice failed ({e}), falling back to built-in renderer")
        else:
            print("LibreOffice not found, using built-in PIL renderer")

        for sheet_name in sheet_names:
            print(f"\n[{sheet_name}] Converting to image...")
            try:
                image = libreoffice_images.get(sheet_name) or self.excel_to_image(sheet_name)
                if image is None:
                    continue
                image_path = self.output_dir / f"{sheet_name}_screenshot.png"
                image.save(image_path)
                self.sheet_images[sheet_name] = image_path
                print(f"[{sheet_name}] Reviewing with LLM...")
                self.sheet_reviews[sheet_name] = self.review_image(image, sheet_name)
                print(f"[{sheet_name}] Done.")
            except Exception as e:
                print(f"[{sheet_name}] Error: {e}")
                self.sheet_reviews[sheet_name] = f"Error: {e}"

    def generate_report(self):
        """Generate a DOCX report with sheet images and review results."""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc = Document()
        doc.add_heading("Excel 审阅报告", level=1)
        doc.add_paragraph(f"文件：{os.path.basename(self.excel_path)}")
        doc.add_paragraph(f"生成时间：{timestamp}")
        doc.add_paragraph(f"审阅工作表数量：{len(self.sheet_reviews)}")
        doc.add_paragraph(
            "说明：本报告由程序自动生成。每个工作表先转为截图，再由模型输出中文 Markdown 审阅意见，并解析写入本报告。"
        )

        for sheet_name in self.sheet_reviews.keys():
            image_path = self.sheet_images.get(sheet_name)
            review = self.sheet_reviews.get(sheet_name, "未生成审阅意见。")
            review = self._sanitize_review_markdown(review)

            doc.add_heading(f"工作表：{sheet_name}", level=2)
            if image_path and Path(image_path).exists():
                doc.add_paragraph("截图预览：")
                doc.add_picture(str(image_path), width=Inches(7.2))
            else:
                doc.add_paragraph("截图预览：未找到截图文件。")

            doc.add_heading("审阅结果", level=3)
            self._append_markdown_to_doc(doc, review)

        report_path = self.output_dir / "review_report.docx"
        doc.save(str(report_path))
        print(f"Report saved: {report_path}")
        return report_path


def main():
    parser = argparse.ArgumentParser(
        description="Convert Excel sheets to images and review with LLM vision"
    )
    parser.add_argument("excel_file", help="Path to the Excel file")
    parser.add_argument("-o", "--output", default="output", help="Output directory (default: output)")
    parser.add_argument("-m", "--model", default=None, help="Model to use (default: OPENAI_MODEL env or gpt-4o)")
    parser.add_argument("-u", "--url", default=None, help="Base URL for OpenAI-compatible API endpoint (default: OPENAI_BASE_URL env)")
    parser.add_argument("-s", "--sheets", nargs="+", default=None, help="Specific sheet names to process")
    parser.add_argument("-n", "--limit", type=int, default=None, help="Limit to first N sheets")
    args = parser.parse_args()

    if not os.path.exists(args.excel_file):
        print(f"Error: File not found: {args.excel_file}")
        sys.exit(1)
    if not args.excel_file.endswith((".xlsx", ".xls")):
        print("Error: File must be an Excel file (.xlsx or .xls)")
        sys.exit(1)

    reviewer = ExcelImageReviewer(args.excel_file, args.output, args.model, args.url)
    reviewer.process_excel(sheets=args.sheets, limit=args.limit)
    report_path = reviewer.generate_report()
    print(f"\nDone! Report: {report_path}")


if __name__ == "__main__":
    main()
